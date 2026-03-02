"""Generate synthetic .docx template files for sample data."""

from pathlib import Path

from docx import Document


SAMPLE_DIR = Path(__file__).parent.parent / "data" / "sample"


def create_nda_v1() -> None:
    doc = Document()
    doc.add_heading("Mutual Non-Disclosure Agreement", level=0)
    doc.add_paragraph(
        "This Mutual Non-Disclosure Agreement (\"Agreement\") is entered into "
        "as of [DATE] by and between Acme Corporation (\"Disclosing Party\") "
        "and [RECEIVING PARTY] (\"Receiving Party\")."
    )
    doc.add_heading("1. Definition of Confidential Information", level=1)
    doc.add_paragraph(
        "\"Confidential Information\" means any non-public information disclosed "
        "by either party to the other party, whether orally, in writing, or by "
        "inspection, including but not limited to business plans, financial data, "
        "technical specifications, and trade secrets."
    )
    doc.add_heading("2. Obligations of Receiving Party", level=1)
    doc.add_paragraph(
        "The Receiving Party shall: (a) hold all Confidential Information in strict "
        "confidence; (b) not disclose Confidential Information to any third party "
        "without prior written consent; (c) use Confidential Information solely "
        "for the purpose of evaluating a potential business relationship."
    )
    doc.add_heading("3. Term", level=1)
    doc.add_paragraph(
        "This Agreement shall remain in effect for a period of two (2) years "
        "from the date of execution."
    )
    doc.add_heading("4. Remedies", level=1)
    doc.add_paragraph(
        "The parties acknowledge that a breach of this Agreement may cause "
        "irreparable harm. Standard legal remedies shall apply."
    )
    doc.add_heading("5. Governing Law", level=1)
    doc.add_paragraph(
        "This Agreement shall be governed by and construed in accordance with "
        "the laws of the State of Delaware."
    )
    doc.save(SAMPLE_DIR / "nda_template_v1.docx")
    print("Created nda_template_v1.docx")


def create_nda_v2() -> None:
    doc = Document()
    doc.add_heading("Mutual Non-Disclosure Agreement", level=0)
    doc.add_paragraph(
        "This Mutual Non-Disclosure Agreement (\"Agreement\") is entered into "
        "as of [DATE] by and between Acme Corporation (\"Disclosing Party\") "
        "and [RECEIVING PARTY] (\"Receiving Party\")."
    )
    doc.add_heading("1. Definition of Confidential Information", level=1)
    doc.add_paragraph(
        "\"Confidential Information\" means any non-public information disclosed "
        "by either party to the other party, whether orally, in writing, or by "
        "inspection, including but not limited to business plans, financial data, "
        "technical specifications, and trade secrets. Confidential Information "
        "expressly excludes any data used for AI model training purposes."
    )
    doc.add_heading("2. Obligations of Receiving Party", level=1)
    doc.add_paragraph(
        "The Receiving Party shall: (a) hold all Confidential Information in strict "
        "confidence; (b) not disclose Confidential Information to any third party "
        "without prior written consent; (c) use Confidential Information solely "
        "for the purpose of evaluating a potential business relationship; "
        "(d) not use Confidential Information to train, fine-tune, or otherwise "
        "improve any artificial intelligence or machine learning model."
    )
    doc.add_heading("3. Term", level=1)
    doc.add_paragraph(
        "This Agreement shall remain in effect for a period of three (3) years "
        "from the date of execution."
    )
    doc.add_heading("4. Remedies", level=1)
    doc.add_paragraph(
        "The parties acknowledge that a breach of this Agreement may cause "
        "irreparable harm. In addition to standard legal remedies, the "
        "non-breaching party shall be entitled to seek injunctive relief."
    )
    doc.add_heading("5. Return of Materials", level=1)
    doc.add_paragraph(
        "Upon termination or expiration of this Agreement, each party shall "
        "promptly return or destroy all Confidential Information received from "
        "the other party, and certify in writing that it has done so."
    )
    doc.add_heading("6. Governing Law", level=1)
    doc.add_paragraph(
        "This Agreement shall be governed by and construed in accordance with "
        "the laws of the State of Delaware."
    )
    doc.save(SAMPLE_DIR / "nda_template_v2.docx")
    print("Created nda_template_v2.docx")


def create_consulting_agreement() -> None:
    doc = Document()
    doc.add_heading("Consulting Services Agreement", level=0)
    doc.add_paragraph(
        "This Consulting Services Agreement (\"Agreement\") is entered into "
        "as of [DATE] by and between Acme Corporation (\"Client\") and "
        "Apex Consulting LLC (\"Consultant\")."
    )
    doc.add_heading("1. Scope of Services", level=1)
    doc.add_paragraph(
        "Consultant shall provide the services described in Exhibit A attached "
        "hereto (the \"Services\"). Consultant shall perform the Services in a "
        "professional and workmanlike manner consistent with industry standards."
    )
    doc.add_heading("2. Compensation", level=1)
    doc.add_paragraph(
        "Client shall pay Consultant at the rate of [RATE] per hour. Payment "
        "shall be made within thirty (30) days of receipt of Consultant's invoice."
    )
    doc.add_heading("3. Intellectual Property", level=1)
    doc.add_paragraph(
        "All work product created by Consultant in the performance of the "
        "Services shall be the sole property of Client. Consultant hereby "
        "assigns all right, title, and interest in such work product to Client."
    )
    doc.add_heading("4. Confidentiality", level=1)
    doc.add_paragraph(
        "Consultant shall hold in confidence all proprietary information "
        "received from Client and shall not disclose such information to "
        "any third party without Client's prior written consent."
    )
    doc.add_heading("5. Term and Termination", level=1)
    doc.add_paragraph(
        "This Agreement shall commence on [START DATE] and continue until "
        "[END DATE], unless terminated earlier. Either party may terminate "
        "this Agreement upon thirty (30) days written notice."
    )
    doc.save(SAMPLE_DIR / "consulting_agreement_v1.docx")
    print("Created consulting_agreement_v1.docx")


def create_software_license() -> None:
    doc = Document()
    doc.add_heading("Software License Agreement", level=0)
    doc.add_paragraph(
        "This Software License Agreement (\"Agreement\") is entered into "
        "as of [DATE] by and between Acme Corporation (\"Licensor\") and "
        "[LICENSEE] (\"Licensee\")."
    )
    doc.add_heading("1. Grant of License", level=1)
    doc.add_paragraph(
        "Licensor hereby grants Licensee a non-exclusive, non-transferable "
        "license to use the software product known as [SOFTWARE NAME] "
        "(the \"Software\") for Licensee's internal business purposes."
    )
    doc.add_heading("2. Restrictions", level=1)
    doc.add_paragraph(
        "Licensee shall not: (a) modify, adapt, or create derivative works "
        "of the Software; (b) reverse engineer, disassemble, or decompile "
        "the Software; (c) sublicense, rent, or lease the Software to any "
        "third party."
    )
    doc.add_heading("3. Service Level Agreement", level=1)
    doc.add_paragraph(
        "Licensor shall maintain 99.9% uptime for the Software, measured "
        "monthly. Scheduled maintenance windows shall not count against "
        "the uptime commitment."
    )
    doc.add_heading("4. Fees", level=1)
    doc.add_paragraph(
        "Licensee shall pay the annual license fee of [AMOUNT] within "
        "thirty (30) days of invoice. Late payments shall accrue interest "
        "at 1.5% per month."
    )
    doc.add_heading("5. Termination", level=1)
    doc.add_paragraph(
        "Either party may terminate this Agreement upon sixty (60) days "
        "written notice. Upon termination, Licensee shall cease all use "
        "of the Software and destroy all copies in its possession."
    )
    doc.save(SAMPLE_DIR / "software_license_v1.docx")
    print("Created software_license_v1.docx")


if __name__ == "__main__":
    SAMPLE_DIR.mkdir(parents=True, exist_ok=True)
    create_nda_v1()
    create_nda_v2()
    create_consulting_agreement()
    create_software_license()
    print("All sample documents created.")
