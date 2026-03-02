import difflib
from io import BytesIO

from docx import Document as DocxDocument


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(BytesIO(file_bytes))
    lines: list[str] = []
    for para in doc.paragraphs:
        lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            lines.append(row_text)
    return "\n".join(lines)


def compute_diff(
    text_a: str,
    text_b: str,
    label_a: str = "Version A",
    label_b: str = "Version B",
) -> dict:
    lines_a = text_a.splitlines(keepends=True)
    lines_b = text_b.splitlines(keepends=True)

    unified = list(difflib.unified_diff(lines_a, lines_b, fromfile=label_a, tofile=label_b))
    unified_text = "".join(unified)

    side_by_side = _compute_side_by_side(lines_a, lines_b)

    return {
        "unified_diff": unified_text,
        "side_by_side": side_by_side,
    }


def _compute_side_by_side(lines_a: list[str], lines_b: list[str]) -> list[dict]:
    matcher = difflib.SequenceMatcher(None, lines_a, lines_b)
    result = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for i, j in zip(range(i1, i2), range(j1, j2)):
                result.append({
                    "line_number_left": i + 1,
                    "line_number_right": j + 1,
                    "content_left": lines_a[i].rstrip("\n"),
                    "content_right": lines_b[j].rstrip("\n"),
                    "change_type": "equal",
                })
        elif tag == "replace":
            left = lines_a[i1:i2]
            right = lines_b[j1:j2]
            max_len = max(len(left), len(right))
            for k in range(max_len):
                result.append({
                    "line_number_left": (i1 + k + 1) if k < len(left) else None,
                    "line_number_right": (j1 + k + 1) if k < len(right) else None,
                    "content_left": left[k].rstrip("\n") if k < len(left) else "",
                    "content_right": right[k].rstrip("\n") if k < len(right) else "",
                    "change_type": "replace",
                })
        elif tag == "delete":
            for i in range(i1, i2):
                result.append({
                    "line_number_left": i + 1,
                    "line_number_right": None,
                    "content_left": lines_a[i].rstrip("\n"),
                    "content_right": "",
                    "change_type": "delete",
                })
        elif tag == "insert":
            for j in range(j1, j2):
                result.append({
                    "line_number_left": None,
                    "line_number_right": j + 1,
                    "content_left": "",
                    "content_right": lines_b[j].rstrip("\n"),
                    "change_type": "insert",
                })
    return result
