from io import BytesIO

from docx import Document

from app.services.diff_service import compute_diff, extract_text_from_docx


def _make_docx(text: str) -> bytes:
    doc = Document()
    doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_text():
    content = _make_docx("Hello World")
    text = extract_text_from_docx(content)
    assert "Hello World" in text


def test_extract_text_with_table():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    buf = BytesIO()
    doc.save(buf)
    text = extract_text_from_docx(buf.getvalue())
    assert "A" in text
    assert "D" in text


def test_compute_diff_equal():
    result = compute_diff("same text", "same text")
    assert result["unified_diff"] == ""
    assert all(l["change_type"] == "equal" for l in result["side_by_side"])


def test_compute_diff_changes():
    result = compute_diff("line one\nline two", "line one\nline three")
    changes = [l for l in result["side_by_side"] if l["change_type"] != "equal"]
    assert len(changes) > 0
    assert result["unified_diff"] != ""


def test_compute_diff_insertion():
    result = compute_diff("line one\nline two", "line one\nline two\nline three")
    non_equal = [l for l in result["side_by_side"] if l["change_type"] != "equal"]
    assert len(non_equal) > 0
    assert any(l["content_right"] == "line three" for l in non_equal)
